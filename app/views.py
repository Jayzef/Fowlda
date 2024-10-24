from django.shortcuts import render, redirect
from django.views import View
from django.contrib import messages
from .forms import ArquivoForm
import PyPDF2
import zipfile
from docx import Document
import io
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pandas as pd
from odf.opendocument import OpenDocumentSpreadsheet
from odf.table import Table, TableRow, TableCell
from odf.text import P
from moviepy.editor import VideoFileClip, AudioFileClip, ImageClip
import os
import tempfile
from PIL import Image

class IndexView(View):
    def get(self, request):
        return render(request, 'index.html')
    
    def post(self, request):
        return render(request, 'index.html')

class ConversaoView(View):
    def get(self, request):
        form = ArquivoForm()
        return render(request, 'conversao.html', {'form': form})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            formato_selecionado = request.POST['formato']  # Captura o formato escolhido
            base_name = uploaded_file.name.rsplit('.', 1)[0]  # Nome base do arquivo sem extensão

            if uploaded_file.name.endswith('.pdf'):
                if formato_selecionado == 'txt':
                    # Converter PDF para TXT
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    txt_content = []

                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        txt_content.append(text)

                    # Retornar o arquivo TXT
                    response = HttpResponse(content_type='text/plain')
                    response['Content-Disposition'] = f'attachment; filename="{base_name}.txt"'
                    response.write("\n".join(txt_content).encode('utf-8'))
                    return response

                elif formato_selecionado == 'docx':
                    # Converter PDF para DOCX
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    doc = Document()
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        doc.add_paragraph(text)

                    # Retornar o arquivo DOCX
                    docx_io = io.BytesIO()
                    doc.save(docx_io)
                    docx_io.seek(0)

                    response = HttpResponse(docx_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename="{base_name}.docx"'
                    return response

            elif uploaded_file.name.endswith('.docx'):
                if formato_selecionado == 'txt':
                    # Converter DOCX para TXT
                    doc = Document(uploaded_file)
                    txt_content = []

                    for paragraph in doc.paragraphs:
                        txt_content.append(paragraph.text)

                    # Retornar o arquivo TXT
                    response = HttpResponse(content_type='text/plain')
                    response['Content-Disposition'] = f'attachment; filename="{base_name}.txt"'
                    response.write("\n".join(txt_content).encode('utf-8'))
                    return response

                elif formato_selecionado == 'pdf':
                    # Converter DOCX para PDF
                    doc = Document(uploaded_file)
                    pdf_io = io.BytesIO()
                    c = canvas.Canvas(pdf_io, pagesize=letter)
                    width, height = letter

                    for paragraph in doc.paragraphs:
                        c.drawString(72, height - 72, paragraph.text)  # Margem de 1 polegada
                        height -= 12  # Mover para baixo

                    c.save()
                    pdf_io.seek(0)

                    # Retornar o arquivo PDF
                    response = HttpResponse(pdf_io.getvalue(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{base_name}.pdf"'
                    return response

            elif uploaded_file.name.endswith('.txt'):
                if formato_selecionado == 'docx':
                    # Converter TXT para DOCX
                    txt_content = uploaded_file.read().decode('utf-8').splitlines()
                    doc = Document()
                    for line in txt_content:
                        doc.add_paragraph(line)

                    # Retornar o arquivo DOCX
                    docx_io = io.BytesIO()
                    doc.save(docx_io)
                    docx_io.seek(0)

                    response = HttpResponse(docx_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename="{base_name}.docx"'
                    return response

                elif formato_selecionado == 'pdf':
                    # Converter TXT para PDF
                    txt_content = uploaded_file.read().decode('utf-8').splitlines()
                    pdf_io = io.BytesIO()
                    c = canvas.Canvas(pdf_io, pagesize=letter)
                    width, height = letter

                    for line in txt_content:
                        c.drawString(72, height - 72, line)  # Margem de 1 polegada
                        height -= 12  # Mover para baixo

                    c.save()
                    pdf_io.seek(0)

                    # Retornar o arquivo PDF
                    response = HttpResponse(pdf_io.getvalue(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{base_name}.pdf"'
                    return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversao')

        return render(request, 'conversao.html', {'form': form})

    
class ConversaoPLAView(View):
    def get(self, request):
        form = ArquivoForm()
        return render(request, 'conversaoPLA.html', {'form': form})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            base_name = uploaded_file.name.rsplit('.', 1)[0]

            if uploaded_file.name.endswith('.ods'):
                # Converter ODS para XLSX
                df = pd.read_excel(uploaded_file, engine='odf')
                xlsx_io = io.BytesIO()
                df.to_excel(xlsx_io, index=False, engine='openpyxl')
                xlsx_io.seek(0)

                response = HttpResponse(xlsx_io, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Disposition'] = f'attachment; filename="{base_name}.xlsx"'
                return response

            elif uploaded_file.name.endswith('.xlsx'):
                # Converter XLSX para ODS
                df = pd.read_excel(uploaded_file, engine='openpyxl')
                ods_io = io.BytesIO()

                # Criar um novo documento ODS
                doc = OpenDocumentSpreadsheet()
                table = Table(name="Sheet1")
                doc.spreadsheet.addElement(table)

                for row in df.itertuples(index=False):
                    table_row = TableRow()
                    for value in row:
                        cell = TableCell()
                        # Adicionar texto como elemento ODF
                        text_element = P(text=str(value))
                        cell.addElement(text_element)
                        table_row.addElement(cell)
                    table.addElement(table_row)

                doc.save(ods_io)
                ods_io.seek(0)

                response = HttpResponse(ods_io, content_type='application/vnd.oasis.opendocument.spreadsheet')
                response['Content-Disposition'] = f'attachment; filename="{base_name}.ods"'
                return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversaoPLA')

        return render(request, 'conversaoPLA.html', {'form': form})
    
class ConversaoVIDView(View):
    def get(self, request):
        form = ArquivoForm()
        return render(request, 'conversaoVID.html', {'form': form})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            base_name = uploaded_file.name.rsplit('.', 1)[0]

            if uploaded_file.name.endswith('.mp4'):
                # Converter MP4 para MP3
                with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as temp_video_file:
                    for chunk in uploaded_file.chunks():
                        temp_video_file.write(chunk)
                    temp_video_file_name = temp_video_file.name

                # Cria o arquivo de áudio temporário
                audio_io = tempfile.NamedTemporaryFile(suffix='.mp3', delete=False)
                
                # Usa um contexto `with` para garantir que o vídeo esteja fechado antes de usar
                with VideoFileClip(temp_video_file_name) as video_clip:
                    audio_clip = video_clip.audio
                    audio_clip.write_audiofile(audio_io.name)
                    audio_clip.close()  # Fecha o clip de áudio

                # Garantindo que o arquivo de áudio seja fechado antes de abrir
                audio_io.close()

                response = HttpResponse(open(audio_io.name, 'rb'), content_type='audio/mpeg')
                response['Content-Disposition'] = f'attachment; filename="{base_name}.mp3"'
                
                # Limpeza de arquivos temporários
                os.remove(temp_video_file_name)
                os.remove(audio_io.name)
                return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversaoVID')

        return render(request, 'conversaoVID.html', {'form': form})
    
class ConversaoIMGView(View):
    def get(self, request):
        form = ArquivoForm()
        return render(request, 'conversaoIMG.html', {'form': form})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            base_name = uploaded_file.name.rsplit('.', 1)[0]

            if uploaded_file.name.endswith('.jpg'):
                # Converter JPG para PNG
                with Image.open(uploaded_file) as img:
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as png_io:
                        img.save(png_io.name, format='PNG')
                        png_io.seek(0)
                        response = HttpResponse(open(png_io.name, 'rb'), content_type='image/png')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.png"'
                
                # Limpeza de arquivos temporários
                os.remove(png_io.name)
                return response

            elif uploaded_file.name.endswith('.png'):
                # Converter PNG para JPG
                with Image.open(uploaded_file) as img:
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as jpg_io:
                        img.convert('RGB').save(jpg_io.name, format='JPEG')
                        jpg_io.seek(0)
                        response = HttpResponse(open(jpg_io.name, 'rb'), content_type='image/jpeg')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.jpg"'

                # Limpeza de arquivos temporários
                os.remove(jpg_io.name)
                return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversaoIMG')

        return render(request, 'conversaoIMG.html', {'form': form})
    
class HistoricoView(View):
    def get(self, request):
        return render(request, 'historico.html')
    
    def post(self, request):
        return render(request, 'historico.html')